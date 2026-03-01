from fastapi import HTTPException
from typing import Callable
from app.enums.format import Format
from pypdf import PdfReader
import fitz
from PIL import Image
import pytesseract
import csv
import docx
import io
from io import StringIO
from io import BytesIO
from odf.opendocument import load
from odf.table import Table, TableRow, TableCell
from odf import text as odf_text
from PIL import Image


Extractor = Callable[[bytes, str | None], tuple[str, dict]]


def extract_content_and_meta(data: bytes, fmt: "Format", filename: str | None) -> tuple[str, dict]:
    if fmt == Format.PDF:
        return extract_pdf_text_and_meta(data)
    if fmt == Format.TXT:
        return extract_txt_text_and_meta(data)
    if fmt == Format.CSV:
        return extract_csv_text_and_meta(data)
    if fmt == Format.XLSX:
        return extract_xlsx_text_and_meta(data)
    if fmt == Format.DOCX:
        return extract_docx_text_and_meta(data)
    if fmt == Format.ODS:
        return extract_ods_text_and_meta(data)
    if fmt == Format.ODT:
        return extract_odt_text_and_meta(data)
    if fmt in (Format.PNG, Format.JPG, Format.JPEG):
        return extract_image_meta_and_optional_ocr(data, filename)
    raise HTTPException(status_code=400, detail=f"Unsupported format: {fmt}")


def extract_pdf_text_and_meta(data: bytes) -> tuple[str, dict]:
    try:
        reader = PdfReader(io.BytesIO(data))
        text = "\n".join((page.extract_text() or "") for page in reader.pages)

        raw_meta = reader.metadata or {}
        def norm_key(k: str) -> str:
            return str(k).lstrip("/").lower()

        meta = {norm_key(k): (str(v) if v is not None else None) for k, v in raw_meta.items()}
        meta.update({"pages": len(reader.pages), "encrypted": bool(reader.is_encrypted)})

        # If no embedded text, try OCR.
        if not text.strip():
            doc = fitz.open(stream=data, filetype="pdf")

            ocr_chunks = []
            for i, page in enumerate(doc):
                pix = page.get_pixmap(dpi=300)

                # Convert pixmap to PNG bytes in-memory.
                png_bytes = pix.tobytes("png")
                img = Image.open(io.BytesIO(png_bytes))

                # OCR.
                ocr_chunks.append(
                    pytesseract.image_to_string(img, lang="eng")
                )

            text = "\n".join(ocr_chunks)

        return text, meta

    except Exception as e:
        raise HTTPException(status_code=400, detail=f"PDF extraction failed: {type(e).__name__}: {e}")

def extract_txt_text_and_meta(data: bytes) -> tuple[str, dict]:
    # best-effort decode
    try:
        text = data.decode("utf-8")
        encoding = "utf-8"
    except UnicodeDecodeError:
        text = data.decode("latin-1", errors="replace")
        encoding = "latin-1"

    meta = {
        "bytes": len(data),
        "encoding": encoding,
        "lines": text.count("\n") + 1 if text else 0,
    }
    return text, meta


def extract_xlsx_text_and_meta(data: bytes) -> tuple[str, dict]:
    try:
        import openpyxl
    except ImportError:
        raise HTTPException(
            status_code=500, detail="openpyxl is not installed (required for .xlsx)")

    from io import BytesIO

    wb = openpyxl.load_workbook(BytesIO(data), data_only=True, read_only=True)
    sheet_names = wb.sheetnames

    parts: list[str] = []
    sheet_stats = {}

    for name in sheet_names:
        ws = wb[name]
        max_row = ws.max_row or 0
        max_col = ws.max_column or 0

        # build TSV-like text
        parts.append(f"=== Sheet: {name} ===")
        non_empty = 0

        for row in ws.iter_rows(values_only=True):
            row_vals = []
            for v in row:
                if v is None:
                    row_vals.append("")
                else:
                    s = str(v)
                    row_vals.append(s)
                    if s != "":
                        non_empty += 1
            # trim trailing empties
            while row_vals and row_vals[-1] == "":
                row_vals.pop()
            if row_vals:
                parts.append("\t".join(row_vals))

        sheet_stats[name] = {
            "rows": max_row,
            "cols": max_col,
            "non_empty_cells_approx": non_empty,
        }

    text = "\n".join(parts)
    meta = {
        "sheets": sheet_names,
        "sheet_stats": sheet_stats,
    }
    return text, meta


def extract_csv_text_and_meta(data: bytes) -> tuple[str, dict]:
    # decode (same logic as TXT)
    try:
        text = data.decode("utf-8")
        encoding = "utf-8"
    except UnicodeDecodeError:
        text = data.decode("latin-1", errors="replace")
        encoding = "latin-1"

    sio = StringIO(text)
    reader = csv.reader(sio)
    rows = list(reader)

    # text representation (TSV-like)
    lines = ["\t".join(map(str, r)) for r in rows if r]
    content = "\n".join(lines)

    meta = {
        "encoding": encoding,
        "rows": len(rows),
        "cols_max": max((len(r) for r in rows), default=0),
    }
    return content, meta


def extract_docx_text_and_meta(data: bytes) -> tuple[str, dict]:
    d = docx.Document(BytesIO(data))
    paragraphs = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    content = "\n".join(paragraphs)

    # basic meta
    meta = {"paragraphs": len(d.paragraphs)}
    try:
        core = d.core_properties
        meta.update({
            "title": core.title,
            "author": core.author,
            "created": core.created.isoformat() if core.created else None,
            "modified": core.modified.isoformat() if core.modified else None,
            "last_modified_by": core.last_modified_by,
        })
    except Exception:
        pass

    return content, meta


def extract_odt_text_and_meta(data: bytes) -> tuple[str, dict]:
    doc = load(BytesIO(data))

    # extract all text nodes
    parts = []
    for node in doc.getElementsByType(odf_text.P):
        # node is paragraph; iterate children
        parts.append(
            "".join((c.data for c in node.childNodes if hasattr(c, "data"))))

    content = "\n".join([p for p in parts if p and p.strip()])

    meta = {}
    try:
        # document meta is optional and varies
        md = doc.meta
        meta["has_meta"] = md is not None
    except Exception:
        pass

    return content, meta


def extract_ods_text_and_meta(data: bytes) -> tuple[str, dict]:
    doc = load(BytesIO(data))

    sheets = doc.spreadsheet.getElementsByType(Table)
    sheet_names = [s.getAttribute("name") for s in sheets]

    parts = []
    sheet_stats = {}

    for s in sheets:
        sname = s.getAttribute("name") or "Sheet"
        parts.append(f"=== Sheet: {sname} ===")
        rows = s.getElementsByType(TableRow)
        max_cols = 0

        for r in rows:
            cells = r.getElementsByType(TableCell)
            max_cols = max(max_cols, len(cells))
            row_vals = []
            for c in cells:
                # concatenate all text:p elements
                ps = c.getElementsByType(odf_text.P)
                cell_text = " ".join(
                    "".join((ch.data for ch in p.childNodes if hasattr(ch, "data"))) for p in ps
                ).strip()
                row_vals.append(cell_text)
            # trim trailing empties
            while row_vals and row_vals[-1] == "":
                row_vals.pop()
            if row_vals:
                parts.append("\t".join(row_vals))

        sheet_stats[sname] = {"rows": len(rows), "cols_max": max_cols}

    content = "\n".join(parts)
    meta = {"sheets": sheet_names, "sheet_stats": sheet_stats}
    return content, meta


def extract_image_meta_and_optional_ocr(data: bytes, filename: str | None) -> tuple[str, dict]:
    img = Image.open(BytesIO(data))

    meta = {
        "format": img.format,
        "mode": img.mode,
        "width": img.size[0],
        "height": img.size[1],
    }

    # EXIF (mostly jpg)
    try:
        exif = img.getexif()
        if exif:
            meta["exif"] = {str(k): str(v) for k, v in exif.items()}
    except Exception:
        pass

    # content: images don't have text unless OCR is used
    content = ""

    # Optional OCR (only if you want it):
    # try:
    #     import pytesseract
    #     content = pytesseract.image_to_string(img)
    #     meta["ocr"] = True
    # except Exception:
    #     meta["ocr"] = False

    return content, meta
